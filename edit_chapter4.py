from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path('/Users/keepwatchthemoon/Desktop/EasyBake咖啡烘焙管理系统_完整修订标注(2).docx')
OUTPUT = Path('/Users/keepwatchthemoon/个人/gitProject/MyCoffeeRoastingBackstage_v1/EasyBake咖啡烘焙管理系统_第四章补充版.docx')
IMAGE_ROOT = Path('/private/var/folders/vh/_z2jj99s7gsd9v1bvd_1snkr0000gn/T/TemporaryItems/com.apple.Photos.NSItemProvider')

IMAGE_PATHS = {
    'IMG_2179.png': IMAGE_ROOT / 'uuid=605B664D-01A8-488C-8771-6C73D11033E8&code=001&library=1&type=1&mode=1&loc=true&cap=true.png' / 'IMG_2179.png',
    'IMG_2180.png': IMAGE_ROOT / 'uuid=924926E4-9F8B-4AA9-AB76-AAE3AFB2583C&code=001&library=1&type=1&mode=1&loc=true&cap=true.png' / 'IMG_2180.png',
    'IMG_2181.png': IMAGE_ROOT / 'uuid=0964F30D-AF55-4922-9CE2-1EA5BB676253&code=001&library=1&type=1&mode=1&loc=true&cap=true.png' / 'IMG_2181.png',
    'IMG_2182.png': IMAGE_ROOT / 'uuid=47C8A483-9D2A-42BD-9058-8C61AB4E7751&code=001&library=1&type=1&mode=1&loc=true&cap=true.png' / 'IMG_2182.png',
    'IMG_2183.png': IMAGE_ROOT / 'uuid=8A2D557C-C332-4180-9D62-7740BC1B705A&code=001&library=1&type=1&mode=1&loc=true&cap=true.png' / 'IMG_2183.png',
    'IMG_2184.png': IMAGE_ROOT / 'uuid=4A9982C0-F5BD-439D-BF52-61584EE77F14&code=001&library=1&type=1&mode=1&loc=true&cap=true.png' / 'IMG_2184.png',
}

RED = RGBColor(0xC0, 0x00, 0x00)


def set_cjk_font(run):
    run.font.name = 'STHeiti'
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'STHeiti')


def add_red_paragraph(document: Document, text: str, bold: bool = False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_cjk_font(run)
    run.font.color.rgb = RED
    run.font.size = Pt(11)
    run.bold = bold
    return paragraph


def add_red_heading(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_cjk_font(run)
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(14)
    return paragraph


def add_feature(document: Document, heading: str, description: str, image_name: str, caption: str):
    document.add_page_break()
    add_red_heading(document, heading)
    add_red_paragraph(document, description)
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_after = Pt(6)
    image_paragraph.add_run().add_picture(str(IMAGE_PATHS[image_name]), width=Inches(3.25))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_cjk_font(caption_run)
    caption_run.font.color.rgb = RED
    caption_run.font.size = Pt(10)


def main():
    for path in IMAGE_PATHS.values():
        if not path.exists():
            raise FileNotFoundError(path)

    document = Document(str(SOURCE))
    add_feature(
        document,
        '4.6 生豆库存页面',
        '【新增】生豆库存页面用于集中查看和维护全部生豆。页面支持按生豆名称、产地、处理法和风味进行搜索，并提供筛选入口；顶部汇总展示总剩余库存和平均采购单价。每张生豆卡片展示名称、豆种、采购成本、默认单份售价、剩余数量及库存进度，用户可通过卡片操作菜单进入详情或执行后续库存管理。',
        'IMG_2179.png',
        '图4.8 生豆库存页面（新增）',
    )
    add_feature(
        document,
        '4.7 生豆详情页面',
        '【新增】生豆详情页面按基础信息、库存与成本、产地与风味等分组展示单支生豆档案。系统记录名称、编号、等级、处理法、豆种、产季、总库存、剩余库存、采购成本、成本模板、默认烘焙量、默认单份售价和默认单份重量，并以标签形式呈现产地、供应商及风味描述，便于后续烘焙计划和成本核算复用。',
        'IMG_2180.png',
        '图4.9 生豆详情页面（新增）',
    )
    add_feature(
        document,
        '4.8 烘焙计划列表页面',
        '【新增】烘焙计划页面用于管理待执行计划和查看烘焙历史。页面提供按计划名称、生豆和烘焙程度搜索的入口，并通过“烘焙计划”“烘焙历史”标签切换视图。计划卡片展示计划名称、总时间、发展时间、生豆、烘焙目标等关键信息，支持从列表进入计划详情或通过操作菜单进行维护。',
        'IMG_2181.png',
        '图4.10 烘焙计划列表页面（新增）',
    )
    add_feature(
        document,
        '4.9 烘焙计划详情页面',
        '【新增】烘焙计划详情页面展示计划概览和按顺序排列的烘焙节点。计划概览包含计划名称、烘豆机型号、批次重量、烘焙目标、用途及当前状态；节点记录时间、事件、操作、炉温、风温、火力和转速等参数，支持完整追溯入豆、回温点、转黄及后续关键操作，为实际烘焙执行和复盘提供依据。',
        'IMG_2182.png',
        '图4.11 烘焙计划详情页面（新增）',
    )
    add_feature(
        document,
        '4.10 财务分析页面',
        '【新增】财务分析页面按全部、本年、本月、本周和今日等时间范围汇总经营指标。页面展示库存预估成本、全部花费、已实现收入、已售出生豆成本、已实现利润、当前库存预估收入、库存预估利润和经营利润，并提供进入明细的操作入口。页面下方展示成本模板及其生豆用量、单价、利润率、包装、能耗等参数，支持经营者快速核对成本与收益。',
        'IMG_2183.png',
        '图4.12 财务分析页面（新增）',
    )
    add_feature(
        document,
        '4.11 系统设置与账户管理页面',
        '【新增】系统设置页面用于管理账户、数据连接、界面外观、烘焙机和 AI 烘焙功能。页面展示当前用户资料、Supabase 连接状态及浅色界面偏好，并提供主动备份、主动上传、注销账号等账户操作；各设置分组可展开查看和维护，为系统运行、数据安全及设备配置提供统一入口。',
        'IMG_2184.png',
        '图4.13 系统设置与账户管理页面（新增）',
    )
    document.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == '__main__':
    main()
